# -*- coding: utf-8 -*-
# app.py — Тестовая система мер ТТР
# Компактный дашборд + карта + чат с ИИ + генерация письма в АНО «Моспром»

import os
import io
import re
import json
from pathlib import Path

import numpy as np
import pandas as pd
import plotly.express as px
import pydeck as pdk
import requests
import streamlit as st
import psycopg2
from pandas import DataFrame
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from docx import Document

# =============================
# Настройки
# =============================
DB_HOST = os.environ.get("PGHOST", "localhost")
DB_PORT = int(os.environ.get("PGPORT", "5433"))
DB_NAME = os.environ.get("PGDATABASE", "Hackaton")
DB_USER = os.environ.get("PGUSER", "postgres")
DB_PASS = os.environ.get("PGPASSWORD", "123")
DB_SCHEMA = "ttr"

OPENAI_BASE = os.environ.get("OPENAI_BASE", "http://26.81.18.206:1234/v1")
CHAT_MODEL  = os.environ.get("CHAT_MODEL", "meta-llama-3.1-8b-instruct")
EMB_MODEL   = os.environ.get("EMB_MODEL", "text-embedding-bge-m3")

INDEX_DIR = Path("knowledge/index")
GEOJSON_PATH = Path(r"c:\Users\dmelnikov\Downloads\world_adm0_custom_exploded.geojson")

st.set_page_config(page_title="Система мер ТТР", layout="wide")
st.title("Система мер ТТР — дашборд")

# =============================
# Соединение с БД
# =============================
def get_conn():
    return psycopg2.connect(
        host=DB_HOST, port=DB_PORT, dbname=DB_NAME, user=DB_USER, password=DB_PASS
    )

# =============================
# Загрузка данных (кэш)
# =============================
@st.cache_data(show_spinner=False)
def load_goods_list() -> pd.DataFrame:
    q = f"SELECT id, hs_code, name FROM {DB_SCHEMA}.goods ORDER BY name"
    with get_conn() as conn:
        return pd.read_sql(q, conn)

@st.cache_data(show_spinner=False)
def load_tariffs(good_id: int) -> DataFrame:
    q = f"SELECT good_id, applied_rate, wto_bound_rate FROM {DB_SCHEMA}.tariffs WHERE good_id = %s"
    with get_conn() as conn:
        return pd.read_sql(q, conn, params=[good_id])

@st.cache_data(show_spinner=False)
def load_series(table: str, good_id: int) -> DataFrame:
    q = f"SELECT year, value_usd_mln FROM {DB_SCHEMA}.{table} WHERE good_id = %s ORDER BY year"
    with get_conn() as conn:
        return pd.read_sql(q, conn, params=[good_id])

@st.cache_data(show_spinner=False)
def load_imports(good_id: int) -> DataFrame:
    q = f"""
    SELECT year, country, 
           COALESCE(value_usd_mln,0) AS value_usd_mln,
           COALESCE(value_tons,0)    AS value_tons,
           country_group
    FROM {DB_SCHEMA}.import_values
    WHERE good_id = %s
    """
    with get_conn() as conn:
        return pd.read_sql(q, conn, params=[good_id])

@st.cache_data(show_spinner=False)
def load_goods_flags(good_id: int) -> DataFrame:
    q = f"SELECT good_id, in_techreg, in_pp1875, in_order4114 FROM {DB_SCHEMA}.goods_flags WHERE good_id = %s"
    with get_conn() as conn:
        return pd.read_sql(q, conn, params=[good_id])

@st.cache_data(show_spinner=False)
def load_geojson(path: Path):
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)

# =============================
# Утилиты расчёта
# =============================
def last_year_present(*dfs):
    years = []
    for df in dfs:
        if df is not None and len(df):
            years.append(int(df["year"].max()))
    return max(years) if years else None

def production_ge_consumption(prod_df: DataFrame, cons_df: DataFrame):
    # сравниваем производство и потребление в последнем общем году
    if not len(prod_df) or not len(cons_df):
        return None, None, None
    common = sorted(set(prod_df["year"]).intersection(set(cons_df["year"])))
    if not common:
        return None, None, None
    y = max(common)
    p = float(prod_df.loc[prod_df["year"]==y, "value_usd_mln"].iloc[0])
    c = float(cons_df.loc[cons_df["year"]==y, "value_usd_mln"].iloc[0])
    return (p >= c), y, (p, c)

def choose_metric(df_year: DataFrame):
    # приоритет стоимости, если нулевая — используем тонны
    if df_year["value_usd_mln"].sum() > 0:
        return "value_usd_mln"
    return "value_tons"

def total_import(df: DataFrame, year: int, metric: str) -> float:
    cur = df[df["year"] == year]
    return float(cur[metric].sum()) if len(cur) else 0.0

def unfriendly_import(df: DataFrame, year: int, metric: str) -> float:
    cur = df[(df["year"] == year) & (df["country_group"] == "unfriendly")]
    return float(cur[metric].sum()) if len(cur) else 0.0

def calc_skc(value_usd_mln: float, value_tons: float):
    # средняя контрактная цена
    if value_tons and value_tons != 0:
        return float(value_usd_mln) / float(value_tons)
    return None

COUNTRY_NAME_ALIASES = {
    "UNITED STATES OF AMERICA": "UNITED STATES",
    "RUSSIAN FEDERATION": "RUSSIA",
    "COTE D'IVOIRE": "CÔTE D'IVOIRE",
    "CZECHIA": "CZECH REPUBLIC",
    "REPUBLIC OF KOREA": "SOUTH KOREA",
    "DEMOCRATIC REPUBLIC OF THE CONGO": "CONGO, DEMOCRATIC REPUBLIC OF THE",
    "CONGO": "CONGO",
    "BOLIVIA (PLURINATIONAL STATE OF)": "BOLIVIA",
    "TANZANIA, UNITED REPUBLIC OF": "TANZANIA",
    "VIET NAM": "VIETNAM",
    "SWAZILAND": "ESWATINI",
    "LAO PEOPLE'S DEMOCRATIC REPUBLIC": "LAOS",
    "PALESTINE, STATE OF": "PALESTINE",
    "NORTH MACEDONIA": "MACEDONIA",
    "MYANMAR": "MYANMAR",
    "BURMA": "MYANMAR",
}
def normalize_country_name(name: str) -> str:
    if not name:
        return ""
    upper = name.strip().upper()
    return COUNTRY_NAME_ALIASES.get(upper, upper)

# =============================
# Выбор нетарифных мер (строго по методике)
# =============================
def non_tariff_analysis(prod_ge_cons, in_tr, in_1875, in_4114):
    measures, notes = [], []
    if prod_ge_cons is False:
        measures.append("Мера 6")
        notes.append("II.1.1: Производство < потребления → Мера 6")
        return measures, notes

    if prod_ge_cons is True:
        if in_1875:
            measures.append("Мера 6")
            notes.append("II.1.2.1: Товар присутствует в ПП №1875 → Мера 6")
        else:
            measures.append("Мера 4")
            notes.append("II.1.2.1: Товара нет в ПП №1875 → Мера 4")

        if in_tr and (not in_4114):
            measures.append("Мера 5")
            notes.append("II.1.2.2: Сертификация действует и нет в Приказе №4114 → Мера 5")
        else:
            measures.append("Мера 6")
            notes.append("II.1.2.2: Условия не выполнены → Мера 6")
    else:
        measures.append("Мера 6")
        notes.append("II: Нет данных по производству/потреблению → Мера 6")

    return measures, notes

def compute_recommendation(tariffs: DataFrame, prod_df: DataFrame, cons_df: DataFrame,
                           imp_df: DataFrame, flags_df: DataFrame):
    # алгоритм выбора меры: раздел I (тарифы) и II (нетарифные)
    applied = float(tariffs["applied_rate"].iloc[0]) if len(tariffs) else 0.0
    wto     = float(tariffs["wto_bound_rate"].iloc[0]) if len(tariffs) else 0.0

    prod_ge_cons, _, _ = production_ge_consumption(prod_df, cons_df)
    ly = last_year_present(imp_df)

    share_ns = 0.0
    delta_ns = 0.0
    metric_used = None
    branch = None

    if ly is not None and len(imp_df):
        cur_year = imp_df[imp_df["year"] == ly]
        metric = choose_metric(cur_year); metric_used = metric
        total_cur = total_import(imp_df, ly, metric)
        ns_cur = unfriendly_import(imp_df, ly, metric)
        if total_cur > 0:
            share_ns = ns_cur / total_cur * 100.0
        ns_prev = unfriendly_import(imp_df, ly - 1, metric)
        delta_ns = ns_cur - ns_prev

    in_tr   = bool(flags_df["in_techreg"].iloc[0]) if len(flags_df) else False
    in_1875 = bool(flags_df["in_pp1875"].iloc[0])  if len(flags_df) else False
    in_4114 = bool(flags_df["in_order4114"].iloc[0]) if len(flags_df) else False

    measures, notes = [], []

    # I. Тарифные меры
    if share_ns >= 30.0 and delta_ns >= 0:
        if prod_ge_cons is True:
            measures.append("Мера 2"); branch = "4.1.1.1"
            notes.append("4.1.1.1: Производство ≥ потреблению → Мера 2 (НС ≥ 30% и не падает)")
        elif prod_ge_cons is False:
            measures.append("Мера 6"); branch = "4.1.1.2"
            notes.append("4.1.1.2: Производство < потребления → Мера 6 (НС ≥ 30% и не падает)")
        else:
            measures.append("Мера 6"); branch = "4.1.1"
            notes.append("4.1.1: Нет данных по P/C → Мера 6")

    elif share_ns < 30.0:
        if wto > applied and (prod_ge_cons is True):
            measures.append("Мера 1"); branch = "4.2.1.1"
            notes.append("4.2.1.1: Bound > Applied и производство ≥ потреблению → Мера 1")
        elif wto > applied and (prod_ge_cons is False):
            measures.append("Мера 6"); branch = "4.2.1.2"
            notes.append("4.2.1.2: Bound > Applied и производство < потребления → Мера 6")
        elif abs(wto - applied) < 1e-12 and (prod_ge_cons is False):
            grew = False
            if ly is not None:
                total_cur = total_import(imp_df, ly, metric_used or "value_usd_mln")
                total_prev = total_import(imp_df, ly - 1, metric_used or "value_usd_mln")
                grew = total_cur > total_prev

            top1_ok = False
            if ly is not None and len(imp_df):
                last = imp_df[imp_df["year"] == ly].copy()
                metric = choose_metric(last)
                top1 = last.sort_values(metric, ascending=False).head(1)
                rest = last.iloc[1:].copy()
                if len(top1) and len(rest):
                    skc_top = calc_skc(float(top1["value_usd_mln"].iloc[0]), float(top1["value_tons"].iloc[0]))
                    skc_others = []
                    for _, r in rest.iterrows():
                        s = calc_skc(float(r["value_usd_mln"]), float(r["value_tons"]))
                        if s is not None: skc_others.append(s)
                    if skc_top is not None and skc_others:
                        top1_ok = skc_top < min(skc_others)

            if grew and top1_ok:
                measures.append("Мера 3"); branch = "4.2.1.3.1.1"
                notes.append("4.2.1.3.1.1: Импорт растёт и СКЦ топ-1 ниже других → Мера 3")
            else:
                measures.append("Мера 6"); branch = "4.2.1.3.1.2"
                notes.append("4.2.1.3.1.2: Условие СКЦ/роста не выполнено → Мера 6")

        elif abs(wto - applied) < 1e-12 and (prod_ge_cons is True):
            nt_measures, nt_notes = non_tariff_analysis(prod_ge_cons, in_tr, in_1875, in_4114)
            measures.extend(nt_measures); branch = "4.2.1.4 → II"
            notes.extend([f"4.2.1.4 → II: {x}" for x in nt_notes])
        else:
            nt_measures, nt_notes = non_tariff_analysis(prod_ge_cons, in_tr, in_1875, in_4114)
            measures.extend(nt_measures); branch = "4.2.1 → II"
            notes.append("4.2.1: Случай вне прямых пунктов (например bound < applied) → II (нетарифные)")
            notes.extend([f"II: {x}" for x in nt_notes])

    else:
        measures.append("Мера 6"); branch = "консервативный выбор при высоких НС и падении объёма"
        notes.append("НС ≥ 30%, но объём из НС снижается; тарифная эскалация не обоснована → Мера 6")

    # убираем дубликаты мер, сохраняя порядок
    seen = set()
    measures = [m for m in measures if not (m in seen or seen.add(m))]

    summary = {
        "last_year": int(ly) if ly is not None else None,
        "share_ns": float(share_ns),
        "delta_ns": float(delta_ns),
        "prod_ge_cons": prod_ge_cons,
        "applied": float(applied), "wto_bound": float(wto),
        "metric_used": metric_used,
        "branch": branch,
        "in_techreg": in_tr, "in_pp1875": in_1875, "in_order4114": in_4114,
        "notes": notes
    }
    return measures, summary

# =============================
# Экспорт (PDF/DOCX/Письмо)
# =============================
def make_pdf_buffer(report_title: str, summary: dict, measures: list):
    """PDF-справка: короткий связный текст + понятное объяснение тарифов в процентах."""
    from textwrap import wrap

    def wrap_lines(s, width=100):
        out = []
        for line in s.split("\n"):
            out.extend(wrap(line, width=width)) if line.strip() else out.append("")
        return out

    def tariff_line(applied: float, wto: float) -> str:
        # Если bound > 0: покажем долю применяемой ставки от bound и «запас»
        if wto and wto > 0:
            applied_pct = (applied / wto) * 100.0
            headroom_pct = max(0.0, 100.0 - applied_pct)
            return (f"Тариф: применяемая ставка {applied:.3f} — это примерно {applied_pct:.1f}% "
                    f"от максимально допустимой; запас до bound ≈ {headroom_pct:.1f}% (ВТО: {wto:.3f}).")
        # bound == 0 → долю посчитать нельзя
        return (f"Тариф: применяемая ставка {applied:.3f}. Максимально допустимый уровень (ВТО) указан как {wto:.3f}, "
                f"поэтому долю в процентах скорректировать невозможно.")

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    width, height = A4
    x, y = 2*cm, height - 2*cm

    def line(s, w=96):
        nonlocal y
        for seg in wrap_lines(s, width=w):
            if y < 2*cm:
                c.showPage(); y = height - 2*cm; c.setFont("Helvetica", 11)
            c.drawString(x, y, seg)
            y -= 0.7*cm

    # Заголовок
    c.setFont("Helvetica-Bold", 14)
    c.drawString(x, y, report_title)
    y -= 1.1*cm

    c.setFont("Helvetica", 11)

    year = summary.get("last_year")
    branch = summary.get("branch") or "—"
    share_ns = summary.get("share_ns")
    delta_ns = summary.get("delta_ns")
    prod_ge_cons = summary.get("prod_ge_cons")
    applied = float(summary.get("applied", 0.0))
    wto = float(summary.get("wto_bound", 0.0))

    # 1) Краткий итог
    line("Краткий итог")
    c.setFont("Helvetica-Bold", 11)
    line(f"Рекомендовано: {', '.join(measures)}.")
    c.setFont("Helvetica", 11)
    if year:
        line(f"Основание: расчёт по методике за {year} год (ветка: {branch}).")
    else:
        line(f"Основание: расчёт по методике (ветка: {branch}).")

    # 2) Ключевые ориентиры периода анализа
    line("")
    line("Ключевые ориентиры периода анализа")
    bullets = []

    # Тариф — с объяснением в процентах
    bullets.append(tariff_line(applied, wto))

    if prod_ge_cons is True:
        bullets.append("Производство в сопоставимом периоде не ниже потребления.")
    elif prod_ge_cons is False:
        bullets.append("Производство в сопоставимом периоде ниже потребления.")
    else:
        bullets.append("Сопоставление производства и потребления учтено в алгоритме.")

    if share_ns is not None and delta_ns is not None and year:
        sign = "+" if delta_ns >= 0 else "−"
        bullets.append(
            f"Импорт за {year}: учитывается доля поставок и её динамика; изменение к предыдущему периоду {sign}{abs(delta_ns):.2f}."
        )

    for b in bullets:
        line(f"• {b}")

    # 3) Обоснование по шагам методики
    notes = summary.get("notes", []) or []
    if notes:
        line("")
        line("Пояснение по шагам методики")
        for n in notes:
            line(f"• {n}")

    # 4) Рекомендации
    line("")
    line("Рекомендации")
    line("• Поддерживать мониторинг показателей и пересчитывать меру при обновлении данных.")
    line("• При необходимости подготовить пакет обоснований для эскалации/коррекции меры через профильные ведомства.")

    c.showPage()
    c.save()
    buf.seek(0)
    return buf


def make_docx_buffer(report_title: str, summary: dict, measures: list):
    """DOCX-справка: нарратив + проценты по тарифам."""
    def tariff_line(applied: float, wto: float) -> str:
        if wto and wto > 0:
            applied_pct = (applied / wto) * 100.0
            headroom_pct = max(0.0, 100.0 - applied_pct)
            return (f"Тариф: применяемая ставка {applied:.3f} — это примерно {applied_pct:.1f}% "
                    f"от максимально допустимой; запас до bound ≈ {headroom_pct:.1f}% (ВТО: {wto:.3f}).")
        return (f"Тариф: применяемая ставка {applied:.3f}. Максимально допустимый уровень (ВТО) указан как {wto:.3f}, "
                f"поэтому долю в процентах скорректировать невозможно.")

    doc = Document()
    doc.add_heading(report_title, level=1)

    year = summary.get("last_year")
    branch = summary.get("branch") or "—"
    share_ns = summary.get("share_ns")
    delta_ns = summary.get("delta_ns")
    prod_ge_cons = summary.get("prod_ge_cons")
    applied = float(summary.get("applied", 0.0))
    wto = float(summary.get("wto_bound", 0.0))

    # Краткий итог
    doc.add_paragraph().add_run("Краткий итог").bold = True
    p = doc.add_paragraph()
    p.add_run(f"Рекомендовано: {', '.join(measures)}. ").bold = True
    if year:
        p.add_run(f"Основание — расчёт по методике за {year} год (ветка: {branch}).")
    else:
        p.add_run(f"Основание — расчёт по методике (ветка: {branch}).")

    # Ключевые ориентиры периода анализа
    doc.add_paragraph().add_run("Ключевые ориентиры периода анализа").bold = True
    doc.add_paragraph(tariff_line(applied, wto), style="List Bullet")

    if prod_ge_cons is True:
        doc.add_paragraph("Производство в сопоставимом периоде не ниже потребления.", style="List Bullet")
    elif prod_ge_cons is False:
        doc.add_paragraph("Производство в сопоставимом периоде ниже потребления.", style="List Bullet")
    else:
        doc.add_paragraph("Сопоставление производства и потребления учтено в алгоритме.", style="List Bullet")

    if share_ns is not None and delta_ns is not None and year:
        sign = "+" if delta_ns >= 0 else "−"
        doc.add_paragraph(
            f"Импорт за {year}: учитывается доля поставок и её динамика; изменение к предыдущему периоду {sign}{abs(delta_ns):.2f}.",
            style="List Bullet"
        )

    # Обоснование по шагам
    notes = summary.get("notes", []) or []
    if notes:
        doc.add_paragraph().add_run("Пояснение по шагам методики").bold = True
        for n in notes:
            doc.add_paragraph(n, style="List Bullet")

    # Рекомендации
    doc.add_paragraph().add_run("Рекомендации").bold = True
    doc.add_paragraph("Поддерживать мониторинг показателей и пересчитывать меру при обновлении данных.", style="List Bullet")
    doc.add_paragraph("При необходимости подготовить пакет обоснований для эскалации/коррекции меры через профильные ведомства.", style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def build_mosprom_letter(good_row, measures, summary, imp_df, tariffs, flags):
    doc = Document()
    doc.add_heading("Обращение в АНО «Моспром»", level=1)
    doc.add_paragraph("От: [Наименование предприятия]")
    doc.add_paragraph(f"Товар: {good_row['name']} (код ТН ВЭД {good_row['hs_code']})")
    doc.add_paragraph(f"Итоговая мера: {', '.join(measures)} (логика: {summary.get('branch') or '—'})")
    if not tariffs.empty:
        doc.add_paragraph(f"Применяемая ставка: {tariffs['applied_rate'].iloc[0]:.3f}; Bound (ВТО): {tariffs['wto_bound_rate'].iloc[0]:.3f}")

    doc.add_heading("Краткое обоснование (автоматически):", level=2)
    for n in summary.get("notes", []):
        doc.add_paragraph(n, style="List Bullet")

    if len(imp_df):
        y = int(imp_df['year'].max())
        doc.add_heading(f"Импорт по странам, {y} (топ-5)", level=2)
        last = imp_df[imp_df["year"]==y].sort_values("value_usd_mln", ascending=False).head(5)
        for _, r in last.iterrows():
            cg = r.get("country_group") or "—"
            doc.add_paragraph(f"{r['country']}: {float(r['value_usd_mln']):.2f} млн $ ({cg})", style="List Bullet")

    doc.add_heading("Просим:", level=2)
    doc.add_paragraph("Содействовать в формировании пакета документов для инициирования рассмотрения вопроса на уровне профильных федеральных органов исполнительной власти.", style="List Bullet")
    doc.add_paragraph("Провести оценку целесообразности инициативы с учётом представленных данных и статистики.", style="List Bullet")

    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf

# =============================
# Сводка товара (для скрытого промпта)
# =============================
def summarize_current_good(good_row, tariffs, prod_df, cons_df, imp_df, measures, summary):
    lines = []
    lines.append(f"Товар: {good_row['name']} (ТН ВЭД {good_row['hs_code']})")
    if not tariffs.empty:
        lines.append(f"Тарифы: применяемая {tariffs['applied_rate'].iloc[0]:.3f}, bound ВТО {tariffs['wto_bound_rate'].iloc[0]:.3f}")
    if len(prod_df) and len(cons_df):
        common = sorted(set(prod_df['year']).intersection(cons_df['year']))
        if common:
            y = max(common)
            p = float(prod_df.loc[prod_df['year']==y, 'value_usd_mln'].iloc[0])
            c = float(cons_df.loc[cons_df['year']==y, 'value_usd_mln'].iloc[0])
            lines.append(f"Производство/Потребление {y}: {p:.2f}/{c:.2f} млн $ (prod_ge_cons={summary.get('prod_ge_cons')})")
    if len(imp_df):
        y = int(imp_df['year'].max())
        total = float(imp_df.loc[imp_df['year']==y, 'value_usd_mln'].sum())
        ns = float(imp_df[(imp_df['year']==y)&(imp_df['country_group']=='unfriendly')]['value_usd_mln'].sum())
        lines.append(f"Импорт {y}: всего {total:.2f} млн $, НС {ns:.2f} млн $ (доля НС {summary.get('share_ns'):.2f}%, delta_ns {summary.get('delta_ns'):.2f})")
    lines.append(f"Итоговые меры: {', '.join(measures)}; логика: {summary.get('branch') or '—'}")
    return "\n".join(lines)

# =============================
# Индекс RAG (если нужен для расширения чата)
# =============================
@st.cache_resource(show_spinner=False)
def load_rag_index():
    rec_path = INDEX_DIR / "records.json"
    emb_path = INDEX_DIR / "embeddings.npy"
    if not rec_path.exists() or not emb_path.exists():
        return None, None
    with open(rec_path, "r", encoding="utf-8") as f:
        records = json.load(f)
    embs = np.load(emb_path)
    norms = np.linalg.norm(embs, axis=1, keepdims=True)
    embs = embs / np.clip(norms, 1e-8, None)
    return records, embs

def embed_query(text: str) -> np.ndarray:
    url = f"{OPENAI_BASE}/embeddings"
    resp = requests.post(url, json={"model": EMB_MODEL, "input": [text]}, timeout=60)
    resp.raise_for_status()
    vec = np.array(resp.json()["data"][0]["embedding"], dtype="float32")
    vec = vec / max(np.linalg.norm(vec), 1e-8)
    return vec

def retrieve(query: str, top_k=5):
    records, embs = load_rag_index()
    if records is None:
        return []
    q = embed_query(query)
    sims = embs @ q
    idx = np.argsort(-sims)[:top_k]
    out = []
    for i in idx:
        r = records[int(i)]
        out.append({"doc": r["doc"], "chunk_id": r["chunk_id"], "text": r["text"], "score": float(sims[i])})
    return out

# =============================
# Чат: помощник и фильтры ответов
# =============================
CHAT_BOX_CSS = """
<style>
.ai-box { background:#000; color:#fff; border:1px solid #333;
  border-radius:12px; padding:14px 16px; margin-top:8px;
  white-space:pre-wrap; line-height:1.45; }
.ai-title { font-weight:700; margin-bottom:6px; font-size:16px; }
</style>
"""
st.markdown(CHAT_BOX_CSS, unsafe_allow_html=True)

SYSTEM_PROMPT = (
    "Ты аналитик по таможенно-тарифному регулированию. "
    "Отвечай уверенно и кратко на основе ПОЛУЧЕННЫХ данных: итоговых мер, логики ветки и заметок алгоритма, "
    "которые я даю в скрытом контексте. "
    "Запрещено придумывать меры, источники или внешние основания, которых нет в контексте. "
    "Не упоминай внутренние механизмы (БД, RAG). "
    "Не используй формулировки «недостаточно данных», «вне текста», «допущение». "
    "Структура ответа: 1) Краткий итог (какая мера и почему — 1–2 предложения, фраза «на основе полученных данных»), "
    "2) Обоснование по шагам методики (Шаг 4.1/4.2/II и подпункты — только из переданных заметок), "
    "3) Рекомендации: до 2 коротких пункта. "
    "Не добавляй разделы «Данные, использованные в расчёте» и не перечисляй сырые числа."
)

FORBIDDEN_PATTERNS = [
    r"(?is)Данные,?\s*использован[ыо]\s*в\s*расч[её]те.*",
    r"(?is)Неопредел[её]нност[ьи].*",
    r"(?is)Допущени[ея].*",
    r"(?is)\bвне\s*текста\b.*",
    r"(?is)согласно\s+данным\s+из\s+БД.*",
    r"(?is)документ\W*\[\d+\].*",
    r"(?is)постановлен[иья]\s+правительств[аи].*",
    r"(?is)указ\s+президента.*",
    r"(?is)publication\.pravo\.gov\.ru.*",
]
def sanitize_ai(text: str) -> str:
    cleaned = text
    for pat in FORBIDDEN_PATTERNS:
        cleaned = re.sub(pat, "", cleaned)
    cleaned = cleaned.strip()
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned

def clamp_measures_in_text(text: str, measures: list[str]) -> str:
    all_tags = re.findall(r"Мера\s*([1-6])", text)
    if not all_tags:
        return text
    allowed = set(m.strip() for m in measures)
    def repl(m):
        tag = f"Мера {m.group(1)}"
        return tag if tag in allowed else ", ".join(allowed) or "Мера 6"
    return re.sub(r"Мера\s*([1-6])", repl, text)

def chat_completion(messages, temperature=0.15, max_tokens=900):
    url = f"{OPENAI_BASE}/chat/completions"
    payload = {"model": CHAT_MODEL, "messages": messages, "temperature": temperature, "max_tokens": max_tokens}
    resp = requests.post(url, json=payload, timeout=120)
    resp.raise_for_status()
    return resp.json()["choices"][0]["message"]["content"]

def make_grounding_message(good_row, measures, summary):
    notes = summary.get("notes", []) or []
    branch = summary.get("branch") or "—"
    ly = summary.get("last_year")
    share_ns = summary.get("share_ns")
    delta_ns = summary.get("delta_ns")
    prod_ge_cons = summary.get("prod_ge_cons")

    lines = []
    lines.append(f"Товар: {good_row['name']} (ТН ВЭД {good_row['hs_code']}).")
    lines.append(f"Итоговые меры алгоритма: {', '.join(measures)}.")
    lines.append(f"Логика ветки: {branch}.")
    if ly is not None:
        lines.append("Контекст последнего периода и сравнений учтён в расчёте алгоритма.")
    if prod_ge_cons is True:
        lines.append("Производство не ниже потребления в последнем сопоставимом периоде.")
    elif prod_ge_cons is False:
        lines.append("Производство ниже потребления в последнем сопоставимом периоде.")
    if share_ns is not None and delta_ns is not None:
        lines.append("Доля и динамика поставок из отдельных групп стран учтены.")
    if notes:
        lines.append("Заметки алгоритма:")
        for n in notes:
            lines.append(f"- {n}")
    return "\n".join(lines)

# =============================
# Данные для страницы
# =============================
goods_df = load_goods_list()
if goods_df.empty:
    st.error("В БД нет товаров. Загрузите справочник ttr.goods.")
    st.stop()

goods_df["display"] = goods_df["name"] + " (" + goods_df["hs_code"] + ")"
selected_good = st.selectbox("Выберите товар:", goods_df["display"])
good_id = int(goods_df.loc[goods_df["display"] == selected_good, "id"].iloc[0])
good_row = goods_df.loc[goods_df["id"] == good_id].iloc[0]

tariffs = load_tariffs(good_id)
prod_df = load_series("production", good_id)
cons_df = load_series("consumption", good_id)
imp_df  = load_imports(good_id)
flags   = load_goods_flags(good_id)

measures, summary = compute_recommendation(tariffs, prod_df, cons_df, imp_df, flags)

# =============================
# Левое меню (вкладки)
# =============================
section = st.sidebar.radio("Разделы", ["Дашборд", "Карта", "Чат", "Выгрузки"])

# =============================
# Дашборд
# =============================
if section == "Дашборд":
    st.markdown(
        f"""
        <div style="
          padding:14px 16px; border-radius:14px;
          background:#000; color:#fff;
          border:1px solid #333; margin: 8px 0 18px 0;">
          <div style="font-size:20px; font-weight:600;">Итоговая мера: {', '.join(measures)}</div>
          <div style="opacity:0.85; font-size:14px; margin-top:4px;">
            Логика выбора: <code style="color:#fff;">{summary.get('branch') or '—'}</code>
          </div>
        </div>
        """,
        unsafe_allow_html=True
    )

    st.markdown("### 📊 Основная информация")
    c1, c2, c3, c4 = st.columns(4)
    with c1:
        st.metric("Код ТН ВЭД", good_row["hs_code"])
    with c2:
        st.metric("Наименование", good_row["name"])
    with c3:
        st.metric("Применяемая ставка", f"{tariffs['applied_rate'].iloc[0]:.3f}" if not tariffs.empty else "—")
    with c4:
        st.metric("Bound (ВТО)", f"{tariffs['wto_bound_rate'].iloc[0]:.3f}" if not tariffs.empty else "—")

    left, right = st.columns([1.2, 1])
    with left:
        st.markdown("#### Производство и потребление")
        if len(prod_df) and len(cons_df):
            ts = prod_df.rename(columns={"value_usd_mln":"Производство, млн $"}).merge(
                cons_df.rename(columns={"value_usd_mln":"Потребление, млн $"}),
                on="year", how="outer"
            ).sort_values("year").rename(columns={"year":"Год"})
            fig_ts = px.line(ts, x="Год", y=["Производство, млн $","Потребление, млн $"], markers=True, title=None)
            st.plotly_chart(fig_ts, use_container_width=True)
        else:
            st.info("Нет данных по производству/потреблению.")

    with right:
        st.markdown("#### Динамика импорта (млн $)")
        if len(imp_df):
            imp_sum = imp_df.groupby("year", as_index=False)["value_usd_mln"].sum().rename(
                columns={"value_usd_mln":"Импорт, млн $", "year":"Год"})
            fig_imp = px.bar(imp_sum, x="Год", y="Импорт, млн $", title=None)
            st.plotly_chart(fig_imp, use_container_width=True)
        else:
            st.info("Нет данных по импорту.")

    with st.expander("📋 Пояснения и параметры расчёта"):
        s = dict(summary)
        if s.get("share_ns") is not None:
            s["share_ns"] = round(s["share_ns"], 2)
        if s.get("delta_ns") is not None:
            s["delta_ns"] = round(s["delta_ns"], 3)
        st.json(s)

# =============================
# Карта
# =============================
elif section == "Карта":
    st.markdown("### 🌍 География импорта")
    if len(imp_df):
        latest_year = int(imp_df["year"].max())
        imp_year = imp_df[imp_df["year"] == latest_year].copy()

        value_by_name = {}
        group_by_name = {}
        for _, r in imp_year.iterrows():
            name_norm = normalize_country_name(str(r["country"]))
            value_by_name[name_norm] = value_by_name.get(name_norm, 0.0) + float(r["value_usd_mln"])
            grp = str(r["country_group"]) if pd.notna(r["country_group"]) else ""
            grp_key = "UNDEFINED"
            if grp == "unfriendly":
                grp_key = "UNFRIENDLY"
            elif grp == "friendly":
                grp_key = "FRIENDLY"
            elif grp:
                grp_key = grp.upper()
            prev = group_by_name.get(name_norm)
            if prev is None or (prev == "FRIENDLY" and grp_key == "UNFRIENDLY"):
                group_by_name[name_norm] = grp_key

        max_val = max(value_by_name.values()) if value_by_name else 1.0

        try:
            geojson = load_geojson(GEOJSON_PATH)
            for feat in geojson["features"]:
                shape_name = normalize_country_name(feat["properties"].get("shapeName", ""))
                val = float(value_by_name.get(shape_name, 0.0))
                grp = group_by_name.get(shape_name, "UNDEFINED")
                feat["properties"]["import_val"] = val
                feat["properties"]["group_label"] = ("Дружественная" if grp == "FRIENDLY"
                                                     else "Недружественная" if grp == "UNFRIENDLY"
                                                     else "—")
                intensity = int(200 * (val / max_val)) if max_val > 0 else 0
                if grp == "UNFRIENDLY":
                    color = [200 + min(55, intensity), 60, 60, 170 if val > 0 else 40]
                elif grp == "FRIENDLY":
                    color = [60, 180 + min(55, intensity), 80, 170 if val > 0 else 40]
                else:
                    color = [180, 180, 180, 50]
                feat["properties"]["fill_color"] = color

            tooltip = {
                "html": "<b>{shapeName}</b>"
                        "<br/>Статус: {group_label}"
                        "<br/>Импорт: {import_val} млн $",
                "style": {"backgroundColor": "rgba(255,255,255,0.92)", "color": "black"}
            }
            fill_layer = pdk.Layer("GeoJsonLayer", geojson, pickable=True, stroked=False,
                                   filled=True, get_fill_color="properties.fill_color",
                                   auto_highlight=True)
            outline_layer = pdk.Layer("GeoJsonLayer", geojson, pickable=False, stroked=True,
                                      filled=False, get_line_color=[90, 90, 90, 220],
                                      line_width_min_pixels=0.7)
            view_state = pdk.ViewState(latitude=25, longitude=15, zoom=1.6, min_zoom=0.5, max_zoom=8)
            deck = pdk.Deck(layers=[fill_layer, outline_layer], initial_view_state=view_state,
                            tooltip=tooltip, map_style=None)
            st.pydeck_chart(deck, use_container_width=True)
            st.caption(f"Данные за {latest_year} год. Наведите на страну для деталей.")
        except Exception as e:
            st.warning(f"Не удалось загрузить GeoJSON. Проверь GEOJSON_PATH. Ошибка: {e}")
    else:
        st.info("Нет данных для карты импорта.")

# =============================
# Чат
# =============================
elif section == "Чат":
    st.markdown("### 🤖 Чат")

    if "chat_history" not in st.session_state:
        st.session_state.chat_history = [
            {"role": "system", "content": SYSTEM_PROMPT}
        ]

    col_q1, col_q2 = st.columns([5,1])
    with col_q1:
        user_q = st.text_input("Ваш вопрос:", "", key="user_q")
    with col_q2:
        ask_clicked = st.button("Спросить ИИ", use_container_width=True)

    grounding = make_grounding_message(good_row, measures, summary)

    if ask_clicked and user_q.strip():
        messages = list(st.session_state.chat_history)
        messages.append({"role": "user", "content": f"Основание для ответа (не показывай как источник):\n{grounding}"})
        messages.append({"role": "user", "content": f"Вопрос: {user_q}"})
        try:
            raw = chat_completion(messages, temperature=0.15, max_tokens=900)
            ans = sanitize_ai(raw)
            ans = clamp_measures_in_text(ans, measures)
            st.session_state.chat_history.append({"role": "user", "content": user_q})
            st.session_state.chat_history.append({"role": "assistant", "content": ans})
        except Exception as e:
            st.error(f"Ошибка запроса к модели: {e}")

    if len(st.session_state.chat_history) > 1:
        for m in st.session_state.chat_history[1:]:
            title = "Вы" if m["role"] == "user" else "Ответ ИИ"
            st.markdown(f"<div class='ai-box'><div class='ai-title'>{title}</div>{m['content']}</div>", unsafe_allow_html=True)

    st.markdown("---")
    st.markdown("### 📄 Сформировать письмо в АНО «Моспром»")
    letter = build_mosprom_letter(good_row, measures, summary, imp_df, tariffs, flags)
    st.download_button(
        "Скачать обращение .docx",
        data=letter,
        file_name=f"Обращение_Моспром_{good_row['hs_code'].replace(' ','')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# =============================
# Выгрузки
# =============================
elif section == "Выгрузки":
    st.markdown("### ⬇️ Выгрузки")
    title = f"Справка по товару: {good_row['name']} ({good_row['hs_code']})"
    colA, colB = st.columns(2)
    with colA:
        pdf = make_pdf_buffer(title, summary, measures)
        st.download_button("⬇️ Скачать PDF",
            data=pdf, file_name=f"spravka_{good_row['hs_code'].replace(' ','')}.pdf",
            mime="application/pdf"
        )
    with colB:
        docx = make_docx_buffer(title, summary, measures)
        st.download_button("⬇️ Скачать DOCX",
            data=docx, file_name=f"spravka_{good_row['hs_code'].replace(' ','')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

# =============================
# Подсказки
# =============================
st.markdown(
    """
    **Примечания**  
    • 4.1 применяется только при `НС ≥ 30%` и `delta_ns ≥ 0`.  
    • 4.2 применяется только при `НС < 30%`.  
    • При `НС ≥ 30%` и снижении объёма из НС — консервативный подход (без эскалации тарифа).
    """
)
