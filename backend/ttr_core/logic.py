# -*- coding: utf-8 -*-
# ttr_core/logic.py — общий модуль логики ТТР: алгоритм, промт, чат, DOCX

import os
import io
import re
import random
from typing import List, Tuple, Optional, Dict, Any

import numpy as np
import pandas as pd
import requests
from docx import Document

# =============================
# Настройки LLM (как в app.py)
# =============================
OPENAI_BASE = os.environ.get("OPENAI_BASE", "http://26.81.18.206:1234/v1")
CHAT_MODEL  = os.environ.get("CHAT_MODEL", "meta-llama-3.1-8b-instruct")

# Глобальные параметры ответа
MAX_WORDS_DEFAULT = 160  # лимит слов на ответ (подправь при необходимости)

# =============================
# Алгоритм расчёта мер
# =============================
def production_ge_consumption(prod_df: pd.DataFrame, cons_df: pd.DataFrame):
    if prod_df is None or cons_df is None or prod_df.empty or cons_df.empty:
        return None, None, None
    common = sorted(set(prod_df["year"]).intersection(set(cons_df["year"])) )
    if not common:
        return None, None, None
    y = max(common)
    p = float(prod_df.loc[prod_df["year"] == y, "value_usd_mln"].iloc[0])
    c = float(cons_df.loc[cons_df["year"] == y, "value_usd_mln"].iloc[0])
    return (p >= c), y, (p, c)

def choose_metric(df_year: pd.DataFrame):
    return "value_usd_mln" if (not df_year.empty and df_year["value_usd_mln"].sum() > 0) else "value_tons"

def total_import(df: pd.DataFrame, year: int, metric: str) -> float:
    cur = df[df["year"] == year]
    return float(cur[metric].sum()) if len(cur) else 0.0

def unfriendly_import(df: pd.DataFrame, year: int, metric: str) -> float:
    cur = df[(df["year"] == year) & (df["country_group"] == "unfriendly")]
    return float(cur[metric].sum()) if len(cur) else 0.0

def calc_skc(value_usd_mln: float, value_tons: float):
    if value_tons and value_tons != 0:
        return float(value_usd_mln) / float(value_tons)
    return None

def non_tariff_analysis(prod_ge_cons, in_tr, in_1875, in_4114):
    measures, notes = [], []
    if prod_ge_cons is False:
        measures.append("Мера 6")
        notes.append("Производство < потребления → Мера 6")
        return measures, notes

    if prod_ge_cons is True:
        if in_1875:
            measures.append("Мера 6")
            notes.append("Товар присутствует в ПП №1875 → Мера 6")
        else:
            measures.append("Мера 4")
            notes.append("Товара нет в ПП №1875 → Мера 4")

        if in_tr and (not in_4114):
            measures.append("Мера 5")
            notes.append("Сертификация действует и нет в Приказе №4114 → Мера 5")
        else:
            measures.append("Мера 6")
            notes.append("Условия сертификации/4114 не выполнены → Мера 6")
    else:
        measures.append("Мера 6")
        notes.append("Нет данных по производству/потреблению → Мера 6")

    return measures, notes

def compute_recommendation(
    tariffs: pd.DataFrame,
    prod_df: pd.DataFrame,
    cons_df: pd.DataFrame,
    imp_df: pd.DataFrame,
    flags_df: pd.DataFrame
) -> Tuple[List[str], Dict[str, Any]]:
    """Раздел I (тарифные) и II (нетарифные) — логика без выдачи первичных данных."""
    applied = float(tariffs["applied_rate"].iloc[0]) if len(tariffs) else 0.0
    wto     = float(tariffs["wto_bound_rate"].iloc[0]) if len(tariffs) else 0.0

    # последний год в импорте
    ly = int(imp_df["year"].max()) if (imp_df is not None and len(imp_df)) else None

    # P/C
    prod_ge_cons, _, _ = production_ge_consumption(prod_df, cons_df)

    # доля НС и динамика
    share_ns = 0.0
    delta_ns = 0.0
    metric_used = None
    branch = None  # внутренний маркер, в текст не выводим

    if ly is not None and len(imp_df):
        cur_year = imp_df[imp_df["year"] == ly]
        metric = choose_metric(cur_year)
        metric_used = metric
        total_cur = total_import(imp_df, ly, metric)
        ns_cur = unfriendly_import(imp_df, ly, metric)
        if total_cur > 0:
            share_ns = ns_cur / total_cur * 100.0
        ns_prev = unfriendly_import(imp_df, ly - 1, metric)
        delta_ns = ns_cur - ns_prev

    # флаги
    in_tr   = bool(flags_df["in_techreg"].iloc[0]) if len(flags_df) else False
    in_1875 = bool(flags_df["in_pp1875"].iloc[0])  if len(flags_df) else False
    in_4114 = bool(flags_df["in_order4114"].iloc[0]) if len(flags_df) else False

    measures, notes = [], []

    # I. Тарифы
    if share_ns >= 30.0 and delta_ns >= 0:
        if prod_ge_cons is True:
            measures.append("Мера 2"); branch = "prod>=cons & NS>=30% & non-decrease"
            notes.append("Производство ≥ потреблению, доля «НС» ≥ 30% и не падает → Мера 2")
        elif prod_ge_cons is False:
            measures.append("Мера 6"); branch = "prod<cons & NS>=30% & non-decrease"
            notes.append("Производство < потребления при доле «НС» ≥ 30% → Мера 6")
        else:
            measures.append("Мера 6"); branch = "no P/C & NS>=30% & non-decrease"
            notes.append("Нет данных по P/C при доле «НС» ≥ 30% → Мера 6")

    elif share_ns < 30.0:
        if wto > applied and (prod_ge_cons is True):
            measures.append("Мера 1"); branch = "bound>applied & prod>=cons"
            notes.append("Bound > Applied и производство ≥ потреблению → Мера 1")
        elif wto > applied and (prod_ge_cons is False):
            measures.append("Мера 6"); branch = "bound>applied & prod<cons"
            notes.append("Bound > Applied и производство < потребления → Мера 6")
        elif abs(wto - applied) < 1e-12 and (prod_ge_cons is False):
            grew = False
            if ly is not None:
                total_cur = total_import(imp_df, ly, metric_used or "value_usd_mln")
                total_prev = total_import(imp_df, ly - 1, metric_used or "value_usd_mln")
                grew = total_cur > total_prev

            top1_ok = False
            if ly is not None and len(imp_df):
                last = imp_df[imp_df["year"] == ly].copy()
                metric = choose_metric(last)
                top1 = last.sort_values(metric, ascending=False).head(1)
                rest = last.iloc[1:].copy()
                if len(top1) and len(rest):
                    skc_top = calc_skc(float(top1["value_usd_mln"].iloc[0]), float(top1["value_tons"].iloc[0]))
                    skc_others = []
                    for _, r in rest.iterrows():
                        s = calc_skc(float(r["value_usd_mln"]), float(r["value_tons"]))
                        if s is not None: skc_others.append(s)
                    if skc_top is not None and skc_others:
                        top1_ok = skc_top < min(skc_others)

            if grew and top1_ok:
                measures.append("Мера 3"); branch = "applied==bound & prod<cons & growth & top1_skc_ok"
                notes.append("Импорт растёт и СКЦ топ-1 ниже других → Мера 3")
            else:
                measures.append("Мера 6"); branch = "applied==bound & prod<cons & (no growth or no skc cond)"
                notes.append("Условие роста/СКЦ не выполнено → Мера 6")

        elif abs(wto - applied) < 1e-12 and (prod_ge_cons is True):
            nt_m, nt_n = non_tariff_analysis(prod_ge_cons, in_tr, in_1875, in_4114)
            measures.extend(nt_m); branch = "applied==bound & prod>=cons → non-tariff"
            notes.extend(nt_n)
        else:
            nt_m, nt_n = non_tariff_analysis(prod_ge_cons, in_tr, in_1875, in_4114)
            measures.extend(nt_m); branch = "other → non-tariff"
            notes.append("Случай вне прямых правил (например bound < applied) → нетарифные варианты")
            notes.extend(nt_n)

    else:
        # НС ≥ 30% и падение (delta_ns < 0): консервативный выбор — Мера 6
        measures.append("Мера 6")
        branch = "NS>=30% & decrease"
        notes.append("Доля «НС» ≥ 30%, но объём снижается; эскалация не обоснована → Мера 6")

    # Убираем дубликаты мер и любые вне диапазона 1–6
    seen = set()
    measures = [m for m in measures if (m in {f"Мера {i}" for i in range(1,7)}) and not (m in seen or seen.add(m))]

    summary = {
        "last_year": int(ly) if ly is not None else None,
        "share_ns": float(share_ns),
        "delta_ns": float(delta_ns),
        "prod_ge_cons": prod_ge_cons,
        "applied": float(applied),
        "wto_bound": float(wto),
        "metric_used": metric_used,
        "branch": branch,      # внутренний маркер (в ответ не показываем)
        "notes": notes
    }
    return measures, summary

# =============================
# Промт, санитизация и чат
# =============================
SYSTEM_PROMPT = (
"""
Ты — эксперт Минпромторга РФ по мерам ТТП.
Пиши естественно, но кратко, без повторов и канцелярита.

Жёсткие требования:
- НЕ ссылайся на внутренние номера веток/пунктов (например, 4.1.1.1 и т.п.).
- НЕ придумывай меры, которых нет: допустимы только «Мера 1», «Мера 2», …, «Мера 6».
- Всегда расшифровывай «НС» как «недружественные страны».
- Не упоминай внутренние механизмы, БД, RAG, «данные вне текста» и т.п.
- Не плодить повторы: один тезис — один раз.
- Отдавай ответ в заданном стиле и в пределах лимита слов.

Структура ответа:
1) Итог — 1–2 фразы.
2) Почему так — 2–4 тезиса.
3) Что изменит решение — 1–3 тезиса.
4) Что делать дальше — 1–4 конкретных шага (практично, прикладно).
"""
)

FORBIDDEN_PATTERNS = [
    r"(?is)Данные,?\s*использован[ыо]\s*в\s*расч[её]те.*",
    r"(?is)Неопредел[её]нност[ьи].*",
    r"(?is)Допущени[ея].*",
    r"(?is)\bвне\s*текста\b.*",
    r"(?is)согласно\s+данным\s+из\s+БД.*",
    r"(?is)документ\W*\[\d+\].*",
    r"(?is)постановлен[иья]\s+правительств[аи].*",
    r"(?is)указ\s+президента.*",
    r"(?is)publication\.pravo\.gov\.ru.*",
    # запреты упоминаний «ветки» и любых 3+уровневых нумераций вроде 4.2.1.3.1
    r"(?is)\bветк[аи]\b.*",
    r"(?<!\d)(\d+(?:\.\d+){2,})(?!\d)",
]

STYLE_PRESETS = [
    "Стиль: телеграм-тезисы. Пиши маркерами, 4–8 пунктов, без длинных вводных.",
    "Стиль: аналитическая записка. 3–5 коротких абзацев по 1–2 предложения.",
    "Стиль: Q&A. Короткие блоки «Итог / Почему / Что изменит / Что делать дальше».",
    "Стиль: один плотный абзац, но без повторов.",
    "Стиль: чек-лист. Короткие строки-действия и причины."
]

def pick_style(style_seed: Optional[int] = None) -> str:
    rnd = random.Random(style_seed) if style_seed is not None else random
    return rnd.choice(STYLE_PRESETS)

def _smart_sentence_trim(text: str, max_words: int) -> str:
    """
    Обрезает по границе предложения, если текст длиннее max_words.
    Если подходящей точки/вопросительного/восклицательного знака нет — обрезает по словам,
    но старается не оставлять «обрыв».
    """
    if not text:
        return text
    words = re.findall(r"\S+", text)
    if len(words) <= max_words:
        return text.strip()
    # предварительное обрезание
    cut = " ".join(words[:max_words])
    # ищем последнюю границу предложения в пределах среза
    m = re.search(r"(?s)^(.+[\.!\?…])[^\.!\?…]*$", cut)
    if m and len(m.group(1).split()) >= int(max_words * 0.6):
        return m.group(1).strip()
    # fallback: аккуратный обрез без многоточия
    return cut.rstrip(",;:—- ").strip()

def trim_to_words(text: str, max_words: int = MAX_WORDS_DEFAULT) -> str:
    # совместимость: используем «умное» обрезание
    return _smart_sentence_trim(text or "", max_words)

def remove_branch_numbers(text: str) -> str:
    # Вырезаем «4.1.1.1», «2.3.4», «4.2.1.3.1.2» и строки про «ветку»
    t = re.sub(r"(?<!\d)(\d+(?:\.\d+){1,})(?!\d)", "", text or "", flags=re.UNICODE)
    t = re.sub(r"(?im)^\s*логика\s*ветки.*$", "", t)
    t = re.sub(r"(?im)^\s*ветка\s*[:\-–].*$", "", t)
    t = re.sub(r"[ \t]{2,}", " ", t)
    t = re.sub(r"\n{3,}", "\n\n", t).strip()
    return t

def sanitize_ai(text: str, max_words: int = MAX_WORDS_DEFAULT) -> str:
    cleaned = text or ""
    for pat in FORBIDDEN_PATTERNS:
        cleaned = re.sub(pat, "", cleaned)
    # Явная расшифровка: «НС» → «НС (недружественные страны)», если сокращение встречается без расшифровки
    if "НС" in cleaned and "недружественные страны" not in cleaned.lower():
        cleaned = re.sub(r"\bНС\b", "НС (недружественные страны)", cleaned)
    cleaned = remove_branch_numbers(cleaned)
    cleaned = cleaned.strip()
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    cleaned = trim_to_words(cleaned, max_words=max_words)
    return cleaned

def clamp_measures_in_text(text: str, measures: List[str]) -> str:
    """
    Оставляет в тексте только меры из разрешённого множества (Мера 1–6).
    Любые «Мера 7/8/…» будут заменены на список фактических мер, либо «Мера 6».
    """
    allowed = {f"Мера {i}" for i in range(1, 7)}
    # сначала убираем любые упоминания неразрешённых мер
    def repl_any(m):
        tag = f"Мера {m.group(1)}"
        return tag if tag in allowed else (", ".join(measures) if measures else "Мера 6")
    text = re.sub(r"Мера\s*([0-9]+)", repl_any, text or "")
    # затем мягко клампим даже разрешённые к финальному списку
    if measures:
        allowed_now = set(measures)
        def repl_only(m):
            tag = f"Мера {m.group(1)}"
            return tag if tag in allowed_now else (", ".join(measures) or "Мера 6")
        text = re.sub(r"Мера\s*([1-6])", repl_only, text)
    return text

def chat_completion(messages, temperature=0.15, max_tokens=900, style_seed: Optional[int] = None) -> str:
    # Лёгкая стохастика параметров для разнообразия ответов
    rnd = random.Random(style_seed) if style_seed is not None else random
    temperature = max(0.1, min(0.8, temperature + rnd.uniform(-0.05, 0.25)))
    top_p = round(rnd.uniform(0.7, 0.95), 2)
    presence_penalty = round(rnd.uniform(0.2, 0.6), 2)
    frequency_penalty = round(rnd.uniform(0.2, 0.7), 2)

    url = f"{OPENAI_BASE}/chat/completions"
    payload = {
        "model": CHAT_MODEL,
        "messages": messages,
        "temperature": temperature,
        "top_p": top_p,
        "presence_penalty": presence_penalty,
        "frequency_penalty": frequency_penalty,
        "max_tokens": max_tokens
    }
    # Если твой сервер не принимает некоторые поля — убери их из payload
    r = requests.post(url, json=payload, timeout=120)
    r.raise_for_status()
    return r.json()["choices"][0]["message"]["content"]

def make_grounding_message(good_row: Dict[str, Any], measures: List[str], summary: Dict[str, Any]) -> str:
    notes = summary.get("notes", []) or []
    ly = summary.get("last_year")
    share_ns = summary.get("share_ns")
    delta_ns = summary.get("delta_ns")
    prod_ge_cons = summary.get("prod_ge_cons")

    lines = []
    lines.append(f"Товар: {good_row['name']} (ТН ВЭД {good_row['hs_code']}).")
    lines.append(f"Итоговые меры алгоритма: {', '.join(measures)}.")
    if ly is not None:
        lines.append("Расчёт выполнен по последнему сопоставимому периоду.")
    if prod_ge_cons is True:
        lines.append("Производство не ниже потребления в последнем сопоставимом периоде.")
    elif prod_ge_cons is False:
        lines.append("Производство ниже потребления в последнем сопоставимом периоде.")
    if share_ns is not None and delta_ns is not None:
        lines.append("«НС» = недружественные страны; учитывались их доля и её динамика.")
    if notes:
        lines.append("Заметки алгоритма (кратко):")
        for n in notes[:6]:
            lines.append("- " + trim_to_words(remove_branch_numbers(n), 28))
    return "\n".join(lines)

def build_messages_for_chat(user_question: str,
                            grounding: str,
                            style_seed: Optional[int] = None,
                            max_words: int = MAX_WORDS_DEFAULT):
    style_hint = pick_style(style_seed)
    sys = SYSTEM_PROMPT + f"\n\n{style_hint}\nЛимит: не более {max_words} слов."
    return [
        {"role": "system", "content": sys},
        {"role": "user",   "content": f"{grounding}\n\nВопрос пользователя: {user_question}"}
    ]

# =============================
# DOCX генераторы
# =============================
def _tariff_line_for_text(applied: float, wto: float) -> str:
    """
    Человеческое объяснение тарифа с процентами.
    """
    if wto and wto > 0:
        applied_pct = (applied / wto) * 100.0
        headroom_pct = max(0.0, 100.0 - applied_pct)
        return (f"Тариф: применяемая ставка {applied:.3f} — это примерно {applied_pct:.1f}% "
                f"от максимально допустимой; запас до предела ≈ {headroom_pct:.1f}% (уровень ВТО: {wto:.3f}).")
    return (f"Тариф: применяемая ставка {applied:.3f}. Максимально допустимый уровень (ВТО) указан как {wto:.3f}, "
            f"поэтому долю в процентах скорректировать невозможно.")

def build_brief_docx(
    good: Dict[str, Any],
    measures: List[str],
    summary: Dict[str, Any],
    tariffs_df: pd.DataFrame,
    imports_df: pd.DataFrame
) -> io.BytesIO:
    """
    «Справка по товару …» — кратко и по делу.
    """
    doc = Document()
    title = f"Справка по товару: {good['name']} ({good['hs_code']})"
    doc.add_heading(title, level=1)

    year   = summary.get("last_year")
    share_ns = summary.get("share_ns")
    delta_ns = summary.get("delta_ns")
    prod_ge_cons = summary.get("prod_ge_cons")
    applied = float(summary.get("applied", tariffs_df["applied_rate"].iloc[0] if not tariffs_df.empty else 0.0))
    wto     = float(summary.get("wto_bound", tariffs_df["wto_bound_rate"].iloc[0] if not tariffs_df.empty else 0.0))

    # Краткий итог
    doc.add_paragraph().add_run("Краткий итог").bold = True
    p = doc.add_paragraph()
    p.add_run(f"Рекомендовано: {', '.join(measures)}. ").bold = True
    if year:
        p.add_run(f"Расчёт по методике за {year} год.")
    else:
        p.add_run("Расчёт по методике по последним доступным значениям.")

    # Ключевые ориентиры периода анализа
    doc.add_paragraph().add_run("Ключевые ориентиры периода анализа").bold = True
    doc.add_paragraph(_tariff_line_for_text(applied, wto), style="List Bullet")

    if prod_ge_cons is True:
        doc.add_paragraph("Производство в сопоставимом периоде не ниже потребления.", style="List Bullet")
    elif prod_ge_cons is False:
        doc.add_paragraph("Производство в сопоставимом периоде ниже потребления.", style="List Bullet")
    else:
        doc.add_paragraph("Сопоставление производства и потребления учтено в расчёте.", style="List Bullet")

    if year is not None and delta_ns is not None:
        sign = "+" if delta_ns >= 0 else "−"
        doc.add_paragraph(
            f"Импорт за {year}: учтены доля поставок из «НС» и её динамика; изменение к предыдущему периоду {sign}{abs(delta_ns):.2f}.",
            style="List Bullet"
        )

    # Пояснение по шагам методики (заметки алгоритма)
    notes = summary.get("notes", []) or []
    if notes:
        doc.add_paragraph().add_run("Пояснение по шагам методики").bold = True
        for n in notes:
            doc.add_paragraph(remove_branch_numbers(n), style="List Bullet")

    # Рекомендации
    doc.add_paragraph().add_run("Рекомендации").bold = True
    doc.add_paragraph("Поддерживать мониторинг показателей и пересчитывать меру при обновлении данных.", style="List Bullet")
    doc.add_paragraph("При необходимости подготовить пакет обоснований для эскалации/коррекции меры через профильные ведомства.", style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def build_mosprom_docx(
    good: Dict[str, Any],
    measures: List[str],
    summary: Dict[str, Any],
    tariffs_df: pd.DataFrame,
    imports_df: pd.DataFrame
) -> io.BytesIO:
    """
    «Обращение в АНО “Моспром”» — кратко и по делу.
    """
    doc = Document()
    doc.add_heading("Обращение в АНО «Моспром»", level=1)
    doc.add_paragraph("От: [Наименование предприятия]")
    doc.add_paragraph(f"Товар: {good['name']} (код ТН ВЭД {good['hs_code']})")
    doc.add_paragraph(f"Итоговая мера: {', '.join(measures)}")
    if tariffs_df is not None and not tariffs_df.empty:
        doc.add_paragraph(
            f"Применяемая ставка: {tariffs_df['applied_rate'].iloc[0]:.3f}; "
            f"Bound (ВТО): {tariffs_df['wto_bound_rate'].iloc[0]:.3f}"
        )

    # Обоснование/заметки алгоритма
    doc.add_heading("Краткое обоснование (автоматически):", level=2)
    for n in summary.get("notes", []) or []:
        doc.add_paragraph(remove_branch_numbers(n), style="List Bullet")

    # Импорт топ-5
    if imports_df is not None and len(imports_df):
        y = int(imports_df["year"].max())
        doc.add_heading(f"Импорт по странам, {y} (топ-5)", level=2)
        last = imports_df[imports_df["year"] == y].sort_values("value_usd_mln", ascending=False).head(5)
        for _, r in last.iterrows():
            cg = r.get("country_group") or "—"
            doc.add_paragraph(f"{r['country']}: {float(r['value_usd_mln']):.2f} млн $ ({cg})", style="List Bullet")

    # Просим
    doc.add_heading("Просим:", level=2)
    doc.add_paragraph("Содействовать в формировании пакета документов для инициирования рассмотрения вопроса на уровне профильных федеральных органов исполнительной власти.", style="List Bullet")
    doc.add_paragraph("Провести оценку целесообразности инициативы с учётом представленных данных и статистики.", style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# =============================
# Сводка (для скрытого промпта)
# =============================
def summarize_current_good(
    good_row: Dict[str, Any],
    tariffs: pd.DataFrame,
    prod_df: pd.DataFrame,
    cons_df: pd.DataFrame,
    imp_df: pd.DataFrame,
    measures: List[str],
    summary: Dict[str, Any]
) -> str:
    lines = []
    lines.append(f"Товар: {good_row['name']} (ТН ВЭД {good_row['hs_code']})")
    if tariffs is not None and not tariffs.empty:
        lines.append(f"Тарифы: применяемая {tariffs['applied_rate'].iloc[0]:.3f}, bound ВТО {tariffs['wto_bound_rate'].iloc[0]:.3f}")
    if prod_df is not None and cons_df is not None and len(prod_df) and len(cons_df):
        common = sorted(set(prod_df['year']).intersection(cons_df['year']))
        if common:
            y = max(common)
            p = float(prod_df.loc[prod_df['year'] == y, 'value_usd_mln'].iloc[0])
            c = float(cons_df.loc[cons_df['year'] == y, 'value_usd_mln'].iloc[0])
            lines.append(f"Производство/Потребление {y}: {p:.2f}/{c:.2f} млн $ (prod_ge_cons={summary.get('prod_ge_cons')})")
    if imp_df is not None and len(imp_df):
        y = int(imp_df['year'].max())
        total = float(imp_df.loc[imp_df['year'] == y, 'value_usd_mln'].sum())
        ns = float(imp_df[(imp_df['year'] == y) & (imp_df['country_group'] == 'unfriendly')]['value_usd_mln'].sum())
        lines.append(f"Импорт {y}: всего {total:.2f} млн $, «НС» {ns:.2f} млн $ (доля «НС» {summary.get('share_ns'):.2f}%, delta_ns {summary.get('delta_ns'):.2f})")
    lines.append(f"Итоговые меры: {', '.join(measures)}")
    return "\n".join(lines)
